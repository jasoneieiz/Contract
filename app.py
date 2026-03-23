from flask import Flask, request, make_response, send_file
import io, os, tempfile, base64, re, unicodedata
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Font registration ─────────────────────────────────────────────────────────
def find_first(paths):
    for p in paths:
        if os.path.exists(p):
            return p
    return None

FONT_R_PATH   = find_first([os.path.join(BASE_DIR,'fonts','ThaiFont.ttf'),
                             '/usr/share/fonts/truetype/freefont/FreeSerif.ttf'])
FONT_B_PATH   = find_first([os.path.join(BASE_DIR,'fonts','ThaiFontBold.ttf'),
                             '/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf'])
FONT_CJK_PATH = find_first([os.path.join(BASE_DIR,'fonts','CJKFont.ttf'),
                             '/usr/share/fonts/opentype/ipafont-gothic/ipag.ttf'])

pdfmetrics.registerFont(TTFont('DocR',   FONT_R_PATH))
pdfmetrics.registerFont(TTFont('DocB',   FONT_B_PATH))
if FONT_CJK_PATH:
    pdfmetrics.registerFont(TTFont('DocCJK', FONT_CJK_PATH))
    HAS_CJK = True
else:
    HAS_CJK = False

print(f'[font] Thai Regular : {FONT_R_PATH}')
print(f'[font] Thai Bold    : {FONT_B_PATH}')
print(f'[font] CJK          : {FONT_CJK_PATH} (loaded={HAS_CJK})')

# Regex to detect CJK characters
CJK_RE = re.compile(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]')

def best_font(text, bold=False):
    """Pick correct font based on text content."""
    if HAS_CJK and CJK_RE.search(text):
        return 'DocCJK'
    return 'DocB' if bold else 'DocR'


# ── Helpers ───────────────────────────────────────────────────────────────────
def safe_filename(name):
    normalized = unicodedata.normalize('NFKD', name)
    ascii_only = normalized.encode('ascii', 'ignore').decode('ascii')
    cleaned    = re.sub(r'[^\w\s\-]', '', ascii_only).strip()
    return cleaned or 'contract'

def send_bytes(buf, mimetype, ascii_name, utf8_name):
    buf.seek(0)
    pct  = ''.join(f'%{b:02X}' for b in utf8_name.encode('utf-8'))
    resp = make_response(buf.read())
    resp.headers['Content-Type']        = mimetype
    resp.headers['Content-Disposition'] = (
        f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{pct}'
    )
    resp.headers['Cache-Control'] = 'no-cache'
    return resp


# ── Routes ────────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return send_file(os.path.join(BASE_DIR, 'index.html'))


@app.route('/generate/pdf', methods=['POST'])
def gen_pdf():
    try:
        data     = request.json or {}
        title    = data.get('title', 'Document')
        body     = data.get('body',  '')
        sig1_b64 = data.get('sig1',  '')
        sig2_b64 = data.get('sig2',  '')
        s1lbl    = data.get('sig1_label', 'Party 1')
        s2lbl    = data.get('sig2_label', 'Party 2')

        W, H = A4
        MG   = 60
        CW   = W - 2 * MG
        buf  = io.BytesIO()
        c    = canvas.Canvas(buf, pagesize=A4)
        y    = H - MG

        def chk():
            nonlocal y
            if y < MG + 80:
                c.showPage()
                y = H - MG

        def put(text, font=None, size=12, align='left', clr=(0, 0, 0)):
            nonlocal y
            if font is None:
                font = best_font(text)
            c.setFillColorRGB(*clr)
            c.setFont(font, size)
            chk()
            if align == 'center':
                c.drawCentredString(W / 2, y, text)
            else:
                c.drawString(MG, y, text)
            y -= size * 1.7

        def wrap(text, size=12):
            nonlocal y
            font = best_font(text)
            c.setFont(font, size)
            lead  = size * 1.7
            # Use fallback font for splitting if CJK
            try:
                parts = simpleSplit(text, font, size, CW) or ['']
            except Exception:
                parts = [text]
            for p in parts:
                chk()
                # Re-detect font per part (mixed lines)
                pf = best_font(p)
                c.setFont(pf, size)
                c.setFillColorRGB(0, 0, 0)
                c.drawString(MG, y, p)
                y -= lead

        # Title + gold rule
        title_font = best_font(title, bold=True)
        if title_font == 'DocCJK':
            title_font = 'DocCJK'  # CJK has no bold variant, use regular
        else:
            title_font = 'DocB'
        put(title, font=title_font, size=16, align='center', clr=(.1, .1, .1))
        y -= 4
        c.setStrokeColorRGB(.76, .63, .25)
        c.setLineWidth(1.2)
        c.line(MG, y, W - MG, y)
        y -= 16
        c.setStrokeColorRGB(0, 0, 0)
        c.setLineWidth(0.5)

        # Body - handle mixed language lines
        for line in body.split('\n'):
            s = line.rstrip()
            if not s:
                y -= 9
            else:
                wrap(s)

        # Signatures
        y -= 24
        chk()
        sig_y = y

        def draw_sig(b64, x):
            if not b64: return
            try:
                raw = base64.b64decode(b64.split(',')[-1])
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                tmp.write(raw); tmp.close()
                c.drawImage(tmp.name, x, sig_y-50, width=140, height=44, mask='auto')
                os.unlink(tmp.name)
            except Exception:
                pass

        draw_sig(sig1_b64, MG)
        draw_sig(sig2_b64, W / 2 + 12)

        c.setStrokeColorRGB(.3, .3, .3)
        c.setLineWidth(0.6)
        c.line(MG,       sig_y-52, MG+168,   sig_y-52)
        c.line(W/2+12,   sig_y-52, W/2+180,  sig_y-52)

        lbl_font = best_font(s1lbl)
        c.setFont(lbl_font, 10)
        c.setFillColorRGB(.35, .35, .35)
        c.drawString(MG+34,  sig_y-65, f'( {s1lbl} )')
        lbl_font2 = best_font(s2lbl)
        c.setFont(lbl_font2, 10)
        c.drawString(W/2+44, sig_y-65, f'( {s2lbl} )')

        c.save()
        return send_bytes(buf, 'application/pdf',
                          safe_filename(title)+'.pdf', title+'.pdf')

    except Exception as e:
        import traceback
        return {'error': str(e), 'trace': traceback.format_exc()}, 500


@app.route('/generate/docx', methods=['POST'])
def gen_docx():
    try:
        data     = request.json or {}
        title    = data.get('title', 'Document')
        body     = data.get('body',  '')
        sig1_b64 = data.get('sig1',  '')
        sig2_b64 = data.get('sig2',  '')
        s1lbl    = data.get('sig1_label', 'Party 1')
        s2lbl    = data.get('sig2_label', 'Party 2')

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(3.0)
            sec.right_margin  = Cm(2.5)

        def docx_font(text):
            """Pick Word font name for text."""
            if CJK_RE.search(text):
                return 'SimSun'       # Built-in CJK font in Word/LibreOffice
            return 'TH Sarabun New'  # Thai + Latin

        # Title
        tp  = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr  = tp.add_run(title)
        tr.bold           = True
        tr.font.size      = Pt(18)
        tr.font.name      = docx_font(title)
        pPr  = tp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), 'C49A1A')
        pBdr.append(bot); pPr.append(pBdr)
        doc.add_paragraph()

        # Body
        for line in body.split('\n'):
            s = line.rstrip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.space_before = Pt(0)
            if not s:
                p.paragraph_format.space_after = Pt(5)
                continue
            run           = p.add_run(s)
            run.font.size = Pt(14)
            run.font.name = docx_font(s)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature table
        tbl = doc.add_table(rows=3, cols=2)
        tbl.style = 'Table Grid'

        def no_border(cell):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB  = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right','insideH','insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcB.append(el)
            tcPr.append(tcB)

        for row in tbl.rows:
            for cell in row.cells:
                no_border(cell)

        for idx, b64 in enumerate([sig1_b64, sig2_b64]):
            if not b64: continue
            try:
                raw = base64.b64decode(b64.split(',')[-1])
                tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                tmp.write(raw); tmp.close()
                tbl.rows[0].cells[idx].paragraphs[0].add_run().add_picture(
                    tmp.name, width=Cm(5))
                os.unlink(tmp.name)
            except Exception:
                pass

        for cell in tbl.rows[1].cells:
            cell.paragraphs[0].add_run('_' * 38).font.size = Pt(12)

        for cell, lbl in zip(tbl.rows[2].cells, [s1lbl, s2lbl]):
            p           = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run           = p.add_run(f'( {lbl} )')
            run.font.size = Pt(11)
            run.font.name = docx_font(lbl)

        buf = io.BytesIO()
        doc.save(buf)
        return send_bytes(
            buf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            safe_filename(title)+'.docx', title+'.docx')

    except Exception as e:
        import traceback
        return {'error': str(e), 'trace': traceback.format_exc()}, 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 7860))
    print(f'[server] http://localhost:{port}')
    app.run(host='0.0.0.0', port=port, debug=False)
